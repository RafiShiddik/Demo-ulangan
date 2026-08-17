import os
import re
import json
import shutil
import random
import string
import time
import urllib.request
import email.utils
from datetime import datetime
import urllib.parse
try:
    import requests
except ImportError:
    requests = None
try:
    import docx
except ImportError:
    docx = None

from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, static_folder=os.path.join(BASE_DIR, 'static'), static_url_path='/static')
app.secret_key = 'bm_exam_secure_secret_key_2026'

@app.route('/smk budi murni 2.jpg')
@app.route('/smk%20budi%20murni%202.jpg')
def serve_logo():
    logo_path = os.path.join(BASE_DIR, 'smk budi murni 2.jpg')
    if os.path.exists(logo_path):
        return send_file(logo_path)
    return '', 404

@app.route('/static/extracted_images/<class_name>/<filename>')
def serve_extracted_images(class_name, filename):
    local_p = os.path.join(BASE_DIR, 'static', 'extracted_images', class_name, filename)
    if os.path.exists(local_p):
        return send_file(local_p)
    tmp_p = os.path.join('/tmp', 'extracted_images', class_name, filename)
    if os.path.exists(tmp_p):
        return send_file(tmp_p)
    return '', 404

# Global configuration and state
CONFIG = {
    'INITIAL_TOKEN': 'MURNI2'  # Default login token
}

# Real-time state of student sessions (helps proctors monitor exams and manage locks)
ACTIVE_STUDENTS = {}

@app.after_request
def add_header(r):
    """Ensure browser never caches exam or login pages to prevent back-button bypasses."""
    r.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    r.headers["Pragma"] = "no-cache"
    r.headers["Expires"] = "0"
    return r

def clean_filename(name):
    """Clean student name to make it safe for file writing and session indexing."""
    cleaned = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    return cleaned.strip('_').lower()

def generate_unique_exam_token():
    """Generates a unique 6-character alphanumeric token for each student."""
    existing_tokens = {s.get('exam_token') for s in ACTIVE_STUDENTS.values() if s.get('exam_token')}
    while True:
        token = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
        if token not in existing_tokens:
            return token


def parse_docx_math(el):
    """Recursively parse Office Math elements from DOCX XML to output clean HTML tags."""
    ns = {
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }
    tag = el.tag.split('}')[-1]
    
    if tag == 't':
        return el.text or ''
    if tag in ('r', 'oMath'):
        return ''.join(parse_docx_math(c) for c in el)
    if tag == 'f':
        num = el.find('m:num', ns)
        den = el.find('m:den', ns)
        num_str = parse_docx_math(num) if num is not None else ''
        den_str = parse_docx_math(den) if den is not None else ''
        return f'<span class="math-fraction"><span class="math-num">{num_str}</span><span class="math-den">{den_str}</span></span>'
    if tag == 'sSup':
        base = el.find('m:e', ns)
        sup = el.find('m:sup', ns)
        base_str = parse_docx_math(base) if base is not None else ''
        sup_str = parse_docx_math(sup) if sup is not None else ''
        return f'{base_str}<sup>{sup_str}</sup>'
    if tag == 'd':
        e = el.find('m:e', ns)
        e_str = parse_docx_math(e) if e is not None else ''
        return f'({e_str})'
    
    return ''.join(parse_docx_math(c) for c in el)

def extract_images_from_docx(doc_path, class_name):
    """Extracts all images from a docx file and saves them to a static directory."""
    if not os.path.exists(doc_path):
        return {}
    try:
        doc = docx.Document(doc_path)
    except Exception:
        return {}
        
    output_dir = os.path.join(BASE_DIR, "static", "extracted_images", class_name)
    try:
        os.makedirs(output_dir, exist_ok=True)
    except Exception:
        output_dir = os.path.join("/tmp", "extracted_images", class_name)
        try:
            os.makedirs(output_dir, exist_ok=True)
        except Exception:
            pass
    
    saved_images = {}
    for rId, part in doc.part.related_parts.items():
        if 'image' in part.content_type:
            ext = 'png'
            if '/' in part.content_type:
                ext = part.content_type.split('/')[1]
            
            img_filename = f"{rId}.{ext}"
            img_path = os.path.join(output_dir, img_filename)
            
            try:
                if not os.path.exists(img_path):
                    with open(img_path, 'wb') as f:
                        f.write(part._blob)
                saved_images[rId] = f"/static/extracted_images/{class_name}/{img_filename}"
            except Exception:
                pass
                
    return saved_images

def get_soal_base_dir():
    """Finds the absolute path to 'soal matematika' directory across local and serverless environments."""
    candidates = [
        os.path.join(BASE_DIR, 'soal matematika'),
        os.path.join(os.path.dirname(BASE_DIR), 'soal matematika'),
        os.path.join(os.getcwd(), 'soal matematika'),
        '/var/task/soal matematika',
        '/var/task/api/soal matematika',
        os.path.join(BASE_DIR, '..', 'soal matematika')
    ]
    for c in candidates:
        if os.path.exists(c) and os.path.isdir(c):
            return c
    return os.path.join(BASE_DIR, 'soal matematika')

def sync_remote_soal_from_cloud():
    """Syncs missing question files from PythonAnywhere cloud (achmadrafi12) to local 'soal matematika/' directory."""
    token = os.environ.get('PYTHONANYWHERE_API_TOKEN', 'd82ad83dee1aab44732ee2eb022cda0b5ab3aec6')
    user = os.environ.get('PYTHONANYWHERE_USER', 'achmadrafi12')
    if not token or not user or requests is None:
        return
    
    headers = {'Authorization': f'Token {token}'}
    base_dir = get_soal_base_dir()
    os.makedirs(base_dir, exist_ok=True)
    
    try:
        url = f'https://www.pythonanywhere.com/api/v0/user/{user}/files/path/home/{user}/soal%20matematika/'
        r = requests.get(url, headers=headers, timeout=4)
        if r.status_code != 200:
            return
        
        data = r.json()
        for k_name, k_info in data.items():
            if isinstance(k_info, dict) and k_info.get('type') == 'directory':
                class_url = f'https://www.pythonanywhere.com/api/v0/user/{user}/files/path/home/{user}/soal%20matematika/{urllib.parse.quote(k_name)}/'
                cr = requests.get(class_url, headers=headers, timeout=4)
                if cr.status_code == 200:
                    c_data = cr.json()
                    for m_name, m_info in c_data.items():
                        if isinstance(m_info, dict) and m_info.get('type') == 'directory':
                            mat_url = f'https://www.pythonanywhere.com/api/v0/user/{user}/files/path/home/{user}/soal%20matematika/{urllib.parse.quote(k_name)}/{urllib.parse.quote(m_name)}/'
                            mr = requests.get(mat_url, headers=headers, timeout=4)
                            if mr.status_code == 200:
                                m_files = mr.json()
                                local_m_dir = os.path.join(base_dir, k_name, m_name)
                                os.makedirs(local_m_dir, exist_ok=True)
                                
                                for f_name, f_info in m_files.items():
                                    if isinstance(f_info, dict) and f_info.get('type') == 'file':
                                        local_f_path = os.path.join(local_m_dir, f_name)
                                        if not os.path.exists(local_f_path):
                                            file_dl_url = f'https://www.pythonanywhere.com/api/v0/user/{user}/files/path/home/{user}/soal%20matematika/{urllib.parse.quote(k_name)}/{urllib.parse.quote(m_name)}/{urllib.parse.quote(f_name)}'
                                            fr = requests.get(file_dl_url, headers=headers, timeout=8)
                                            if fr.status_code == 200:
                                                with open(local_f_path, 'wb') as f_out:
                                                    f_out.write(fr.content)
                                                print(f"[Cloud Soal Sync] Downloaded {k_name}/{m_name}/{f_name}")
    except Exception as e:
        print(f"[Cloud Soal Sync Error] {e}")

def scan_soal_directory():
    """Scans 'soal matematika/' and returns a dict mapping kelas -> list of available materis."""
    try:
        sync_remote_soal_from_cloud()
    except Exception:
        pass
    base_dir = get_soal_base_dir()
    materi_map = {}
    if not os.path.exists(base_dir):
        return materi_map
        
    for item in sorted(os.listdir(base_dir)):
        class_path = os.path.join(base_dir, item)
        if os.path.isdir(class_path):
            kelas_key = item.replace('Kelas ', '').replace('kelas ', '').strip()
            materis = []
            for sub in sorted(os.listdir(class_path)):
                sub_path = os.path.join(class_path, sub)
                if os.path.isdir(sub_path):
                    materis.append(sub)
            if not materis:
                docx_files = [f for f in os.listdir(class_path) if f.endswith('.docx')]
                if docx_files:
                    materis.append('Matematika Umum')
            if materis:
                materi_map[kelas_key] = materis
            
    return materi_map

def find_materi_folder(kelas, materi=None):
    """Finds the path to the specified class and materi folder."""
    base_dir = get_soal_base_dir()
    if not os.path.exists(base_dir):
        return None
    target_class_dir = None
    norm_k = kelas.replace('Kelas ', '').replace('kelas ', '').strip().upper()
    for item in os.listdir(base_dir):
        if os.path.isdir(os.path.join(base_dir, item)):
            k_key = item.replace('Kelas ', '').replace('kelas ', '').strip().upper()
            if k_key == norm_k:
                target_class_dir = os.path.join(base_dir, item)
                break
                
    if not target_class_dir:
        return None
        
    if materi:
        materi_dir = os.path.join(target_class_dir, materi)
        if os.path.exists(materi_dir) and os.path.isdir(materi_dir):
            return materi_dir
        for sub in os.listdir(target_class_dir):
            sub_path = os.path.join(target_class_dir, sub)
            if os.path.isdir(sub_path) and sub.strip().lower() == materi.strip().lower():
                return sub_path
            
    subs = [os.path.join(target_class_dir, s) for s in os.listdir(target_class_dir) if os.path.isdir(os.path.join(target_class_dir, s))]
    if subs:
        return subs[0]
        
    return target_class_dir

def load_questions(kelas, materi=None):
    """Load multiple choice questions dynamically from DOCX file using robust universal parsing."""
    folder_path = find_materi_folder(kelas, materi)
    if not folder_path or not os.path.exists(folder_path):
        return []

    pg_path = None
    for f in os.listdir(folder_path):
        if f.endswith('.docx') and 'kunci' not in f.lower() and 'essay' not in f.lower() and 'esai' not in f.lower():
            pg_path = os.path.join(folder_path, f)
            break
            
    if not pg_path:
        for f in os.listdir(folder_path):
            if f.endswith('.docx') and 'kunci' not in f.lower():
                pg_path = os.path.join(folder_path, f)
                break
                
    if not pg_path or not os.path.exists(pg_path):
        return []

    images_map = extract_images_from_docx(pg_path, kelas)
    doc = docx.Document(pg_path)
    paragraphs_text = []
    
    for p in doc.paragraphs:
        parts = [parse_docx_math(child) for child in p._element]
        p_text = ''.join(parts).strip()
        
        pPr = p._element.pPr
        has_numPr = pPr is not None and pPr.numPr is not None
        
        blips = []
        def find_blips(el):
            tag = el.tag.split('}')[-1]
            if tag == 'blip':
                for k, v in el.attrib.items():
                    if k.endswith('embed'):
                        blips.append(v)
            for child in el:
                find_blips(child)
        find_blips(p._element)
        
        img_htmls = []
        for rId in blips:
            if rId in images_map:
                img_htmls.append(f'<div class="exam-image-container"><img class="exam-image" src="{images_map[rId]}" alt="Gambar Soal"></div>')
                
        img_str = "".join(img_htmls)
        
        if p_text or img_str:
            if p_text.lower().startswith('peringatan') or p_text.lower() in ['pilihan ganda', 'soal pilihan ganda']:
                continue
                
            combined = (p_text + "<br>" + img_str).strip("<br>").strip() if p_text else img_str
            
            # An image-only paragraph is ONLY a continuation if it does NOT have numPr!
            if not p_text and img_str and not has_numPr and paragraphs_text:
                paragraphs_text[-1] = paragraphs_text[-1] + "<br>" + img_str
            else:
                paragraphs_text.append(combined)
            
    def clean_opt(txt):
        txt = txt.strip()
        txt = re.sub(r'^[a-eA-E][\.\)]\s*', '', txt)
        return txt

    def strip_html(txt):
        return re.sub(r'<[^>]+>', '', txt).strip()

    def is_question_header(p_text):
        clean_text = strip_html(p_text).strip()
        if not clean_text or len(clean_text) < 5:
            return False
        lower_text = clean_text.lower()
        
        if lower_text.startswith(('peringatan', 'pilihan ganda', 'soal pilihan ganda', 'apa itu', 'tipe-tipe', 'translasi', 'teranslasi', 'pergeseran', 'jika a', 'jika b', 'jika c', 'jika k', 'refleksi', 'dilatasi', 'rotasi', 'perhatikan')):
            return False

        if re.match(r'^\s*(?:soal\s*)?\d+[\.\)]', clean_text, re.IGNORECASE):
            return True

        is_q_phrase = (
            clean_text.endswith('?') or 
            clean_text.endswith('...') or 
            'adalah...' in lower_text or 
            'adalah?' in lower_text or
            lower_text.startswith('dibawah ini') or
            lower_text.startswith('berikut ini') or
            lower_text.startswith('titik') or
            lower_text.startswith('sebuah') or
            lower_text.startswith('bayangan') or
            lower_text.startswith('persamaan')
        )
        return is_q_phrase

    questions = []
    letters = ['A', 'B', 'C', 'D', 'E']
    i = 0

    while i < len(paragraphs_text):
        p = paragraphs_text[i]
        if is_question_header(p):
            q_text = re.sub(r'^\s*(?:soal\s*)?\d+[\.\)]\s*', '', p)
            opts = []
            j = i + 1
            while j < len(paragraphs_text) and len(opts) < 5:
                if is_question_header(paragraphs_text[j]):
                    break
                opts.append(paragraphs_text[j])
                j += 1
                
            if len(opts) >= 3:
                choices = {}
                for l_idx, o in enumerate(opts):
                    choices[letters[l_idx]] = clean_opt(o)
                questions.append({
                    'index': len(questions) + 1,
                    'question': q_text,
                    'choices': choices
                })
                i = j
                continue
        i += 1

    return questions

def load_answers(kelas, materi=None):
    """Load answer keys dynamically from DOCX file based on class and materi."""
    folder_path = find_materi_folder(kelas, materi)
    if not folder_path or not os.path.exists(folder_path):
        return {}
        
    key_path = None
    for f in os.listdir(folder_path):
        if 'kunci' in f.lower() and 'essay' not in f.lower() and 'esai' not in f.lower():
            key_path = os.path.join(folder_path, f)
            break
            
    if not key_path:
        for f in os.listdir(folder_path):
            if 'kunci' in f.lower():
                key_path = os.path.join(folder_path, f)
                break
                
    if not key_path or not os.path.exists(key_path):
        return {}
        
    doc = docx.Document(key_path)
    answers = {}
    q_index = 1
    
    for p in doc.paragraphs:
        txt = p.text.strip().upper()
        if not txt or 'KUNCI' in txt or 'JAWABAN' in txt:
            continue
        match = re.search(r'[A-E]', txt)
        if match:
            answers[q_index] = match.group(0)
            q_index += 1
            
    return answers

def load_essay_questions(kelas, materi=None):
    """Load essay questions dynamically from DOCX file if available (guaranteed 5 questions)."""
    folder_path = find_materi_folder(kelas, materi)
    if not folder_path or not os.path.exists(folder_path):
        return []
        
    essay_path = None
    for f in os.listdir(folder_path):
        if ('essay' in f.lower() or 'esai' in f.lower()) and 'kunci' not in f.lower() and f.endswith('.docx'):
            essay_path = os.path.join(folder_path, f)
            break
            
    if not essay_path or not os.path.exists(essay_path):
        return []
        
    doc = docx.Document(essay_path)
    raw_paragraphs = []
    for p in doc.paragraphs:
        parts = [parse_docx_math(child) for child in p._element]
        t = ''.join(parts).strip()
        if t:
            t_low = t.lower()
            if t_low.startswith('peringatan') or 'harap kumpulkan' in t_low or t_low == 'essay !' or t_low == 'essay' or 'isikan caranya' in t_low or 'snbt' in t_low:
                continue
            raw_paragraphs.append(t)
            
    if kelas.upper() == 'XII':
        essay_groups = [
            [0, 1],
            [2],
            [3],
            [4],
            [5]
        ]
        essay_questions = []
        for idx, grp in enumerate(essay_groups):
            q_text = "<br>".join([raw_paragraphs[i] for i in grp if i < len(raw_paragraphs)])
            essay_questions.append({
                'index': idx + 1,
                'question': re.sub(r'^\d+[\.\)]\s*', '', q_text)
            })
        return essay_questions

    questions = []
    for p in raw_paragraphs:
        clean_p = re.sub(r'^\d+[\.\)]\s*', '', p)
        p_lower = clean_p.lower()
        
        is_sub = False
        if questions:
            if p_lower.startswith('tentukan') or p_lower.startswith('hitunglah') or p_lower.startswith('jika g(x)') or re.match(r'^[a-z][\.\)]\s*', p_lower):
                is_sub = True
                
        if is_sub and questions:
            questions[-1]['text'] += '<br>' + clean_p
        else:
            questions.append({
                'text': clean_p
            })
            
    essay_questions = []
    for idx, q in enumerate(questions):
        essay_questions.append({
            'index': idx + 1,
            'question': q['text']
        })
    return essay_questions


def get_shuffled_questions_for_student(student_state):
    """
    Given student state, returns the list of questions for their class and materi,
    with questions ordered according to student_state['question_order'],
    and choices for each question shuffled according to student_state['shuffled_choices'].
    """
    kelas = student_state.get('kelas', 'XI')
    materi = student_state.get('materi', '')
    questions_all = load_questions(kelas, materi)
    q_map = {q['index']: q for q in questions_all}
    
    shuffled_questions = []
    student_order = student_state.get('question_order', [])
    shuffled_choices = student_state.get('shuffled_choices', {})
    
    for idx in student_order:
        if idx not in q_map:
            continue
        q = q_map[idx]
        orig_choices = q['choices']
        
        if str(idx) in shuffled_choices or idx in shuffled_choices:
            mapping = shuffled_choices.get(idx) or shuffled_choices.get(str(idx))
            new_choices = {}
            for disp_letter in sorted(mapping.keys()):
                orig_letter = mapping[disp_letter]
                new_choices[disp_letter] = orig_choices.get(orig_letter, '')
        else:
            new_choices = orig_choices
            
        shuffled_questions.append({
            'index': idx,
            'question': q['question'],
            'choices': new_choices
        })
        
    return shuffled_questions



@app.route('/')
def index():
    if 'siswa' in session:
        status = session.get('status', 'konfirmasi')
        if status == 'aktif':
            return redirect(url_for('ujian'))
        elif status == 'selesai':
            return redirect(url_for('selesai'))
        else:
            return redirect(url_for('konfirmasi'))
    return redirect(url_for('login'))

def check_jurusan_allowed(folder_path, student_jurusan):
    """
    Checks if student_jurusan is allowed to access the material in folder_path based on metadata.json.
    Returns (is_allowed: bool, target_jurusan: str).
    """
    if not folder_path or not os.path.exists(folder_path):
        return True, "Semua Jurusan"
    
    meta_path = os.path.join(folder_path, 'metadata.json')
    if not os.path.exists(meta_path):
        return True, "Semua Jurusan"
        
    try:
        with open(meta_path, 'r', encoding='utf-8') as f:
            meta = json.load(f)
            
        target_jurusan = meta.get('jurusan', 'Semua Jurusan').strip()
        if not target_jurusan or target_jurusan.lower() == 'semua jurusan':
            return True, target_jurusan
            
        if not student_jurusan:
            return True, target_jurusan

        student_j = student_jurusan.strip().upper()
        base_code = student_j.split()[0] if student_j else student_j
        
        allowed_items = [j.strip().upper() for j in target_jurusan.split(',')]
        
        for allowed in allowed_items:
            allowed_clean = allowed.strip()
            if allowed_clean in student_j or student_j in allowed_clean or allowed_clean in base_code or base_code in allowed_clean:
                return True, target_jurusan
            if (base_code in ['TKJ'] and 'TKJ' in allowed_clean) or \
               (base_code in ['MP', 'MPLB'] and ('MP' in allowed_clean or 'MPLB' in allowed_clean)) or \
               (base_code in ['AK', 'AKL'] and ('AK' in allowed_clean or 'AKL' in allowed_clean)) or \
               (base_code in ['BR', 'BD', 'BDP'] and ('BR' in allowed_clean or 'BD' in allowed_clean)) or \
               (base_code in ['DKV'] and 'DKV' in allowed_clean):
                return True, target_jurusan
                
        return False, target_jurusan
    except Exception as e:
        print(f"[Check Jurusan Error] {e}")
        return True, "Semua Jurusan"

@app.route('/api/materis/<kelas>')
def api_materis(kelas):
    student_jurusan = request.args.get('jurusan', '').strip()
    materi_map = scan_soal_directory()
    norm_kelas = kelas.replace('Kelas ', '').replace('kelas ', '').strip().upper()
    materis = []
    for k, v in materi_map.items():
        if k.upper() == norm_kelas:
            for mat in v:
                folder_path = find_materi_folder(kelas, mat)
                is_allowed, _ = check_jurusan_allowed(folder_path, student_jurusan)
                if is_allowed:
                    materis.append(mat)
            break
    if not materis and not student_jurusan:
        materis = ['Matematika Umum']
    return jsonify({'kelas': kelas, 'materis': materis})


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        nama = request.form.get('nama', '').strip()
        kelas = request.form.get('kelas', '').strip()
        jurusan = request.form.get('jurusan', '').strip()
        materi = request.form.get('materi', '').strip()
        
        if not nama or not kelas or not jurusan:
            return render_template('login.html', error='Semua kolom wajib diisi!')
            
        materi_map = scan_soal_directory()
        norm_k = kelas.replace('Kelas ', '').replace('kelas ', '').strip().upper()
        avail_materis = []
        for k, v in materi_map.items():
            if k.upper() == norm_k:
                avail_materis = v
                break
                
        nama_clean = clean_filename(nama)
        
        if nama_clean in ACTIVE_STUDENTS:
            student_state = ACTIVE_STUDENTS[nama_clean]
            
            if student_state['status'] == 'selesai':
                return render_template('login.html', error='Anda telah menyelesaikan ujian ini!')
                
            if student_state['status'] == 'terkunci':
                session['siswa'] = student_state['nama']
                session['kelas'] = student_state['kelas']
                session['jurusan'] = student_state['jurusan']
                session['materi'] = student_state['materi']
                session['nama_clean'] = nama_clean
                session['status'] = 'terkunci'
                return redirect(url_for('konfirmasi'))
            
            if student_state['status'] == 'aktif':
                session['siswa'] = student_state['nama']
                session['kelas'] = student_state['kelas']
                session['jurusan'] = student_state['jurusan']
                session['materi'] = student_state['materi']
                session['nama_clean'] = nama_clean
                session['status'] = 'aktif'
                return redirect(url_for('ujian'))
                
            if student_state['status'] == 'konfirmasi':
                session['siswa'] = student_state['nama']
                session['kelas'] = student_state['kelas']
                session['jurusan'] = student_state['jurusan']
                session['materi'] = student_state['materi']
                session['nama_clean'] = nama_clean
                session['status'] = 'konfirmasi'
                return redirect(url_for('konfirmasi'))

        folder_path = find_materi_folder(kelas, materi)
        if folder_path:
            is_allowed, target_j = check_jurusan_allowed(folder_path, jurusan)
            if not is_allowed:
                return render_template('login.html', error=f'Soal "{materi}" khusus ditujukan untuk jurusan ({target_j}). Jurusan {jurusan} tidak dapat mengikuti ujian ini!')

        questions_all = load_questions(kelas, materi)
        if not questions_all and avail_materis:
            materi = avail_materis[0]
            questions_all = load_questions(kelas, materi)

        if not questions_all:
            return render_template('login.html', error=f'Soal untuk kelas {kelas} ({materi or "Umum"}) belum tersedia!')
            
        q_order = [q['index'] for q in questions_all]
        norm_k = str(kelas).upper().replace('KELAS', '').strip()
        if norm_k != 'XII':
            random.shuffle(q_order)
        
        shuffled_choices = {}
        for q in questions_all:
            q_idx = q['index']
            orig_keys = [k for k, v in q['choices'].items() if v]
            # Keep choice options in original order (A, B, C, D, E) to prevent option mismatch
            mapping = {}
            display_letters = ['A', 'B', 'C', 'D', 'E'][:len(orig_keys)]
            for i, disp_char in enumerate(display_letters):
                mapping[disp_char] = orig_keys[i]
            shuffled_choices[q_idx] = mapping
            
        exam_token = generate_unique_exam_token()
        
        ACTIVE_STUDENTS[nama_clean] = {
            'nama': nama,
            'kelas': kelas,
            'jurusan': jurusan,
            'materi': materi,
            'status': 'konfirmasi',
            'exam_token': exam_token,
            'token_baru': '',
            'lock_time': None,
            'score': None,
            'answers': {},
            'essay_answers': {},
            'lock_count': 0,
            'question_order': q_order,
            'shuffled_choices': shuffled_choices
        }
        
        session['siswa'] = nama
        session['kelas'] = kelas
        session['jurusan'] = jurusan
        session['materi'] = materi
        session['nama_clean'] = nama_clean
        session['status'] = 'konfirmasi'
        
        return redirect(url_for('konfirmasi'))
        
    return render_template('login.html')

@app.route('/konfirmasi', methods=['GET', 'POST'])
def konfirmasi():
    if 'siswa' not in session or 'nama_clean' not in session:
        return redirect(url_for('login'))
        
    nama_clean = session['nama_clean']
    if nama_clean not in ACTIVE_STUDENTS:
        return redirect(url_for('login'))
        
    student_state = ACTIVE_STUDENTS[nama_clean]
    
    if student_state['status'] == 'aktif':
        session['status'] = 'aktif'
        return redirect(url_for('ujian'))

    if request.method == 'POST':
        token_input = request.form.get('token', '').strip().upper()
        if not token_input:
            return render_template('konfirmasi.html', 
                                   nama=session['siswa'], 
                                   kelas=session['kelas'], 
                                   jurusan=session['jurusan'], 
                                   materi=session.get('materi', ''), 
                                   error='Token wajib diisi!')
                                   
        if student_state['status'] == 'terkunci':
            expected_token = student_state.get('token_baru', '')
            if token_input == expected_token:
                ACTIVE_STUDENTS[nama_clean]['status'] = 'aktif'
                ACTIVE_STUDENTS[nama_clean]['token_baru'] = ''
                ACTIVE_STUDENTS[nama_clean]['lock_time'] = None
                
                filepath = f"data_token/{nama_clean}.txt"
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except Exception:
                        pass
                        
                session['status'] = 'aktif'
                return redirect(url_for('ujian'))
            else:
                return render_template('konfirmasi.html', 
                                       nama=session['siswa'], 
                                       kelas=session['kelas'], 
                                       jurusan=session['jurusan'], 
                                       materi=session.get('materi', ''), 
                                       error='Token buka kunci dari pengawas Anda salah atau tidak valid!')
        else:
            expected_token = student_state.get('exam_token', '')
            if token_input == expected_token:
                ACTIVE_STUDENTS[nama_clean]['status'] = 'aktif'
                session['status'] = 'aktif'
                return redirect(url_for('ujian'))
            else:
                return render_template('konfirmasi.html', 
                                       nama=session['siswa'], 
                                       kelas=session['kelas'], 
                                       jurusan=session['jurusan'], 
                                       materi=session.get('materi', ''), 
                                       error='Token Ujian salah! Silakan tanyakan token unik Anda kepada pengawas.')
                                       
    if student_state['status'] == 'selesai':
        return redirect(url_for('selesai'))
    elif student_state['status'] == 'aktif':
        return redirect(url_for('ujian'))
        
    error_msg = None
    if student_state['status'] == 'terkunci':
        error_msg = 'Sesi Anda terkunci karena meninggalkan halaman ujian. Silakan masukkan token buka kunci dari pengawas.'
        
    return render_template('konfirmasi.html', 
                           nama=session['siswa'], 
                           kelas=session['kelas'], 
                           jurusan=session['jurusan'], 
                           materi=session.get('materi', ''), 
                           error=error_msg)

@app.route('/ujian')
def ujian():
    if 'siswa' not in session or 'nama_clean' not in session:
        return redirect(url_for('login'))
        
    nama_clean = session['nama_clean']
    kelas = session['kelas']
    materi = session.get('materi', '')
    
    if nama_clean in ACTIVE_STUDENTS:
        status = ACTIVE_STUDENTS[nama_clean]['status']
        session['status'] = status
    else:
        return redirect(url_for('login'))
        
    if status == 'konfirmasi':
        return redirect(url_for('konfirmasi'))
        
    if status == 'terkunci':
        return redirect(url_for('konfirmasi'))
        
    if status == 'selesai':
        return redirect(url_for('selesai'))
        
    student_state = ACTIVE_STUDENTS[nama_clean]
    questions = get_shuffled_questions_for_student(student_state)
    essay_questions = load_essay_questions(kelas, materi)
    saved_answers = student_state.get('answers', {})
    saved_essay_answers = student_state.get('essay_answers', {})
    lock_count = student_state.get('lock_count', 0)
    
    return render_template('ujian.html', 
                           questions=questions, 
                           essay_questions=essay_questions, 
                           materi=materi, 
                           status=status,
                           saved_answers=saved_answers,
                           saved_essay_answers=saved_essay_answers,
                           lock_count=lock_count)


@app.route('/api/save_answers', methods=['POST'])
def save_answers():
    if 'nama_clean' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    nama_clean = session['nama_clean']
    if nama_clean in ACTIVE_STUDENTS:
        try:
            data = request.get_json(silent=True) or {}
            if 'answers' in data and isinstance(data['answers'], dict):
                ACTIVE_STUDENTS[nama_clean]['answers'].update(data['answers'])
            if 'essay_answers' in data and isinstance(data['essay_answers'], dict):
                ACTIVE_STUDENTS[nama_clean]['essay_answers'].update(data['essay_answers'])
            return jsonify({'status': 'success'})
        except Exception as e:
            return jsonify({'error': str(e)}), 400
    return jsonify({'error': 'Student state not found'}), 404


@app.route('/lock', methods=['POST'])
def lock():
    if 'nama_clean' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    nama_clean = session['nama_clean']
    
    if nama_clean in ACTIVE_STUDENTS and ACTIVE_STUDENTS[nama_clean]['status'] != 'selesai':
        ACTIVE_STUDENTS[nama_clean]['status'] = 'terkunci'
        ACTIVE_STUDENTS[nama_clean]['lock_count'] = ACTIVE_STUDENTS[nama_clean].get('lock_count', 0) + 1
        
        # Save draft answers if included in payload
        try:
            data = request.get_json(silent=True) or {}
            if 'answers' in data and isinstance(data['answers'], dict):
                ACTIVE_STUDENTS[nama_clean]['answers'].update(data['answers'])
            if 'essay_answers' in data and isinstance(data['essay_answers'], dict):
                ACTIVE_STUDENTS[nama_clean]['essay_answers'].update(data['essay_answers'])
        except Exception:
            pass
        
        token_baru = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
        ACTIVE_STUDENTS[nama_clean]['token_baru'] = token_baru
        ACTIVE_STUDENTS[nama_clean]['lock_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        filename = None
        try:
            token_dir = os.path.join(BASE_DIR, 'data_token')
            os.makedirs(token_dir, exist_ok=True)
            filename = os.path.join(token_dir, f"{nama_clean}.txt")
            
            content = (
                "=============================\n"
                "TOKEN BARU SISWA\n"
                "=============================\n"
                f"Nama        : {ACTIVE_STUDENTS[nama_clean]['nama']}\n"
                f"Kelas       : {ACTIVE_STUDENTS[nama_clean]['kelas']}\n"
                f"Materi      : {ACTIVE_STUDENTS[nama_clean].get('materi', '-')}\n"
                f"Jurusan     : {ACTIVE_STUDENTS[nama_clean]['jurusan']}\n"
                f"Token       : {token_baru}\n"
                f"Waktu       : {ACTIVE_STUDENTS[nama_clean]['lock_time']}\n"
                f"Lock Count  : {ACTIVE_STUDENTS[nama_clean]['lock_count']} kali keluar halaman\n"
            )
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
        except Exception as e:
            print(f"[Lock File Save Warning] Could not save token file: {e}")
            
        session['status'] = 'terkunci'
            
        return jsonify({'status': 'locked', 'token_file': filename or '', 'lock_count': ACTIVE_STUDENTS[nama_clean]['lock_count']})
        
    return jsonify({'status': 'ignored'})

@app.route('/unlock', methods=['POST'])
def unlock():
    if 'nama_clean' not in session:
        return redirect(url_for('login'))
        
    nama_clean = session['nama_clean']
    token_input = request.form.get('token_input', '').strip().upper()
    
    if nama_clean in ACTIVE_STUDENTS:
        if ACTIVE_STUDENTS[nama_clean]['status'] == 'aktif':
            session['status'] = 'aktif'
            return redirect(url_for('ujian'))
            
        expected_token = ACTIVE_STUDENTS[nama_clean].get('token_baru', '')
        if token_input and token_input == expected_token:
            ACTIVE_STUDENTS[nama_clean]['status'] = 'aktif'
            ACTIVE_STUDENTS[nama_clean]['token_baru'] = ''
            ACTIVE_STUDENTS[nama_clean]['lock_time'] = None
            
            session['status'] = 'aktif'
            
            filepath = os.path.join(BASE_DIR, "data_token", f"{nama_clean}.txt")
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except Exception:
                    pass
                    
            return redirect(url_for('ujian'))
            
    student_state = ACTIVE_STUDENTS.get(nama_clean, {})
    questions = get_shuffled_questions_for_student(student_state) if student_state else []
    essay_questions = load_essay_questions(session.get('kelas', ''), session.get('materi', ''))
    return render_template('ujian.html', 
                           questions=questions, 
                           essay_questions=essay_questions,
                           materi=session.get('materi', ''),
                           status='terkunci', 
                           error='Token salah! Silakan periksa kembali atau minta ulang pengawas.')

@app.route('/student/status')
def student_status():
    if 'nama_clean' not in session:
        return jsonify({'status': 'none'})
    nama_clean = session['nama_clean']
    if nama_clean in ACTIVE_STUDENTS:
        return jsonify({'status': ACTIVE_STUDENTS[nama_clean]['status']})
    
def get_gspread_client():
    """Initializes gspread client using local JSON file or environment variable with automatic clock-skew compensation."""
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        
        scopes = [
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
        
        # Clock skew compensation for Windows local dev environments
        try:
            res = urllib.request.urlopen('https://www.google.com', timeout=3)
            gtime = email.utils.mktime_tz(email.utils.parsedate_tz(res.headers['Date']))
            drift = int(time.time() - gtime)
        except Exception:
            drift = 0

        if drift > 5 or drift < -5:
            def patched_assertion(self):
                now = int(time.time() - drift)
                payload = {
                    'iss': self._service_account_email,
                    'scope': ' '.join(self._scopes),
                    'aud': self._token_uri,
                    'iat': now - 10,
                    'exp': now + 3600,
                }
                if self._subject:
                    payload['sub'] = self._subject
                payload.update(self._additional_claims)
                import google.auth.jwt
                return google.auth.jwt.encode(self._signer, payload)
            Credentials._make_authorization_grant_assertion = patched_assertion
            
        # 1. Check environment variable GOOGLE_CREDENTIALS_JSON (Render / Cloud / Vercel deployment)
        json_env = os.environ.get('GOOGLE_CREDENTIALS_JSON')
        creds = None
        if json_env:
            import json
            creds_dict = json.loads(json_env)
            if 'private_key' in creds_dict and isinstance(creds_dict['private_key'], str):
                creds_dict['private_key'] = creds_dict['private_key'].replace('\\n', '\n')
            creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
            
        # 2. Check local 'key/' directory or root
        if not creds:
            key_dirs = [os.path.join(BASE_DIR, 'key'), BASE_DIR]
            for kd in key_dirs:
                if os.path.exists(kd):
                    json_files = [os.path.join(kd, f) for f in os.listdir(kd) if f.endswith('.json') and f != 'package.json']
                    if json_files:
                        creds = Credentials.from_service_account_file(json_files[0], scopes=scopes)
                        break
                        
        if creds:
            return gspread.authorize(creds)
    except Exception as e:
        print(f"[Google Sheets Init Error] {e}")
        
    return None


def save_result_to_google_sheet(nama_siswa, kelas, jurusan, materi, score, correct_count, total_count, details, essay_details, lock_count=0):
    """Appends exam result to Google Spreadsheet in a worksheet tab specific to student's class with columns for each question."""
    try:
        gc = get_gspread_client()
        if not gc:
            print("[Google Sheets] Service account credentials not found. Skipping sheet sync.")
            return False
            
        sheet_id = os.environ.get('GOOGLE_SHEET_ID')
        sheet_name = os.environ.get('GOOGLE_SHEET_NAME', 'ulangan harian pertama')
        
        sh = None
        if sheet_id:
            try:
                sh = gc.open_by_key(sheet_id)
            except Exception as e:
                print(f"[Google Sheets] Could not open sheet by ID: {e}")
                
        if not sh:
            try:
                sh = gc.open(sheet_name)
            except Exception as oe:
                print(f"[Google Sheets] Could not open sheet by name '{sheet_name}': {oe}")
                try:
                    sh = gc.create(sheet_name)
                except Exception as ce:
                    print(f"[Google Sheets] Could not create sheet: {ce}")
                    return False
                    
        # Worksheet per class and jurusan (e.g. "Kelas XI - TKJ", "Kelas X - DKV")
        norm_kelas = kelas.strip()
        if not norm_kelas.lower().startswith('kelas'):
            norm_kelas = f"Kelas {norm_kelas}"
        else:
            norm_kelas = norm_kelas.title()

        norm_jurusan = jurusan.strip() if jurusan else ''
        if norm_jurusan and norm_jurusan.lower() != 'semua jurusan':
            target_sheet_title = f"{norm_kelas} - {norm_jurusan}"
        else:
            target_sheet_title = norm_kelas

        try:
            worksheet = sh.worksheet(target_sheet_title)
        except Exception:
            try:
                worksheet = sh.add_worksheet(title=target_sheet_title, rows=100, cols=50)
            except Exception as e:
                print(f"[Google Sheets] Could not create worksheet '{target_sheet_title}': {e}")
                worksheet = sh.sheet1
        
        # Check existing headers
        try:
            existing_rows = worksheet.get_all_values()
        except Exception:
            existing_rows = []
            
        base_headers = [
            "Timestamp", "Nama Siswa", "Kelas", "Jurusan", "Materi",
            "Skor PG (Maks 40)", "Nilai Esai (Manual - Maks 60)", "Total Nilai (100)",
            "Benar PG", "Frekuensi Keluar Halaman (Lock Count)", "Catatan PG"
        ]
        
        # Include question texts in headers for randomized questions transparency
        pg_headers = []
        for d in details:
            q_text = d.get('question', '').strip()
            if len(q_text) > 75:
                q_snippet = q_text[:72] + "..."
            else:
                q_snippet = q_text
            pg_headers.append(f"[PG {d['index']}] {q_snippet}" if q_snippet else f"Soal PG {d['index']}")
            
        essay_headers = []
        for ed in essay_details:
            eq_text = ed.get('question', '').strip()
            if len(eq_text) > 75:
                eq_snippet = eq_text[:72] + "..."
            else:
                eq_snippet = eq_text
            essay_headers.append(f"[ESAI {ed['index']}] {eq_snippet}" if eq_snippet else f"Soal Esai {ed['index']}")
        
        full_headers = base_headers + pg_headers + essay_headers

        if not existing_rows or len(existing_rows) == 0:
            worksheet.append_row(full_headers)
        else:
            current_header_row = existing_rows[0]
            if len(full_headers) >= len(current_header_row):
                try:
                    worksheet.update(values=[full_headers], range_name='A1')
                except Exception:
                    pass

        base_data = [
            datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            nama_siswa,
            kelas,
            jurusan,
            materi,
            score,
            "",
            score,
            f"{correct_count} / {total_count}",
            f"{lock_count} kali",
            "Catatan: Ini adalah nilai sementara Pilihan Ganda. Nilai akhir dapat berubah setelah soal esai diperiksa."
        ]
        
        pg_data = []
        for d in details:
            ans = d.get('student_answer', '') or '-'
            key = d.get('correct_answer', '')
            if d.get('is_correct'):
                pg_data.append(f"{ans} (BENAR)")
            elif ans != '-' and key:
                pg_data.append(f"{ans} (SALAH | Kunci: {key})")
            elif ans == '-':
                pg_data.append(f"- (TIDAK DIISI | Kunci: {key})")
            else:
                pg_data.append(ans)
                
        essay_data = []
        for ed in essay_details:
            ans_e = ed.get('student_answer', '') or '(Tidak diisi)'
            essay_data.append(ans_e)

        row_data = base_data + pg_data + essay_data
        
        worksheet.append_row(row_data)
        print(f"[Google Sheets Success] Recorded exam result for '{nama_siswa}' to worksheet '{target_sheet_title}' in '{sh.title}'")
        return True
    except Exception as e:
        print(f"[Google Sheets Error] Failed to save result: {e}")
        return False


def send_result_to_website_guru(nama_siswa, kelas, jurusan, materi, score, correct_count, total_count, details, essay_details, lock_count=0):
    """Syncs student exam submission to Website Guru (PythonAnywhere) via REST API."""
    token = os.environ.get('PYTHONANYWHERE_API_TOKEN', 'd82ad83dee1aab44732ee2eb022cda0b5ab3aec6')
    
    urls = [
        os.environ.get('WEBSITE_GURU_URL', 'https://achmadrafi12.pythonanywhere.com/api/receive_results'),
        'https://achmadrafi12.pythonanywhere.com/api/upload_hasil',
        'https://achmadrafi12.pythonanywhere.com/api/hasil_ujian',
        'https://achmadrafi12.pythonanywhere.com/api/submit_hasil'
    ]
    
    headers = {
        'Authorization': f'Token {token}',
        'Content-Type': 'application/json'
    }
    
    payload = {
        'nama_siswa': nama_siswa,
        'nama': nama_siswa,
        'kelas': kelas,
        'jurusan': jurusan,
        'materi': materi,
        'score': score,
        'nilai': score,
        'correct_count': correct_count,
        'total_count': total_count,
        'lock_count': lock_count,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'details': details,
        'essay_details': essay_details
    }
    
    data = json.dumps(payload).encode('utf-8')
    for guru_url in urls:
        try:
            req = urllib.request.Request(guru_url, data=data, headers=headers, method='POST')
            with urllib.request.urlopen(req, timeout=5) as resp:
                res_body = resp.read().decode('utf-8')
                print(f"[Website Guru Sync Success] Sent to {guru_url}. Response: {res_body}")
                return True
        except Exception as e:
            print(f"[Website Guru Sync Attempt] Failed for {guru_url}: {e}")
            
    return False



@app.route('/submit', methods=['POST'])
def submit():
    if 'siswa' not in session:
        return redirect(url_for('login'))
        
    nama_clean = session['nama_clean']
    kelas = session['kelas']
    materi = session.get('materi', '')
    jurusan = session['jurusan']
    nama_siswa = session['siswa']
    
    if nama_clean in ACTIVE_STUDENTS and ACTIVE_STUDENTS[nama_clean]['status'] == 'terkunci':
        session.pop('siswa', None)
        session.pop('nama_clean', None)
        session.pop('status', None)
        return redirect(url_for('login', error='Sesi Anda terkunci karena meninggalkan halaman ujian secara otomatis. Silakan minta kode token baru kepada pengawas kelas untuk masuk kembali.'))
        
    if nama_clean in ACTIVE_STUDENTS and ACTIVE_STUDENTS[nama_clean]['status'] == 'selesai':
        return redirect(url_for('selesai'))
        
    questions = load_questions(kelas, materi)
    correct_answers = load_answers(kelas, materi)
    essay_questions = load_essay_questions(kelas, materi)
    
    student_answers = {}
    correct_count = 0
    total_count = len(questions)
    
    details = []
    
    for q in questions:
        q_idx = q['index']
        ans = request.form.get(f'q{q_idx}', '').strip().upper()
        
        if nama_clean in ACTIVE_STUDENTS:
            mapping = ACTIVE_STUDENTS[nama_clean].get('shuffled_choices', {}).get(q_idx, {})
            real_ans = mapping.get(ans, ans)
        else:
            real_ans = ans
            
        student_answers[q_idx] = real_ans
        
        expected = correct_answers.get(q_idx, '')
        is_correct = (real_ans == expected)
        if is_correct:
            correct_count += 1
            
        details.append({
            'index': q_idx,
            'question': q['question'],
            'student_answer': real_ans,
            'correct_answer': expected,
            'is_correct': is_correct,
            'choices': q['choices']
        })
        
    score = round((correct_count / total_count * 40), 2) if total_count > 0 else 0.0
    if isinstance(score, float) and score.is_integer():
        score = int(score)
    
    # Process Essay answers (not auto-graded)
    essay_details = []
    student_essay_answers = {}
    for eq in essay_questions:
        eq_idx = eq['index']
        ans_essay = request.form.get(f'essay{eq_idx}', '').strip()
        student_essay_answers[eq_idx] = ans_essay
        essay_details.append({
            'index': eq_idx,
            'question': eq['question'],
            'student_answer': ans_essay
        })
    
    lock_count = ACTIVE_STUDENTS[nama_clean].get('lock_count', 0) if nama_clean in ACTIVE_STUDENTS else 0

    if nama_clean in ACTIVE_STUDENTS:
        ACTIVE_STUDENTS[nama_clean]['status'] = 'selesai'
        ACTIVE_STUDENTS[nama_clean]['score'] = score
        ACTIVE_STUDENTS[nama_clean]['answers'] = student_answers
        ACTIVE_STUDENTS[nama_clean]['essay_answers'] = student_essay_answers
        
    session['status'] = 'selesai'
    session['score'] = score
    session['essay_answers'] = student_essay_answers
    session['lock_count'] = lock_count
    
    # Save to Google Sheets
    save_result_to_google_sheet(nama_siswa, kelas, jurusan, materi, score, correct_count, total_count, details, essay_details, lock_count)
    
    # Sync result to Website Guru (PythonAnywhere)
    send_result_to_website_guru(nama_siswa, kelas, jurusan, materi, score, correct_count, total_count, details, essay_details, lock_count)
    
    try:
        if not kelas.lower().startswith('kelas'):
            folder_kelas = f"Kelas {kelas}"
        else:
            folder_kelas = kelas[0].upper() + kelas[1:]
            
        if materi:
            result_dir = os.path.join(BASE_DIR, 'hasil ujian', folder_kelas, materi, jurusan, nama_siswa)
        else:
            result_dir = os.path.join(BASE_DIR, 'hasil ujian', folder_kelas, jurusan, nama_siswa)
            
        os.makedirs(result_dir, exist_ok=True)
        
        txt_path = os.path.join(result_dir, 'hasil.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("==================================================\n")
            f.write("HASIL UJIAN SISWA - SMK BUDI MURNI 2\n")
            f.write("==================================================\n")
            f.write(f"Nama                 : {nama_siswa}\n")
            f.write(f"Kelas                : {kelas}\n")
            f.write(f"Materi               : {materi}\n")
            f.write(f"Jurusan              : {jurusan}\n")
            f.write(f"Tanggal              : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Frekuensi Keluar Tab : {lock_count} kali\n")
            f.write(f"Skor PG              : {score} / 40 (Bobot 40%)\n")
            f.write(f"Benar PG             : {correct_count} dari {total_count} soal\n")
            f.write("--------------------------------------------------\n")
            f.write("CATATAN NILAI:\n")
            f.write("Catatan: Ini adalah nilai sementara Pilihan Ganda. Nilai akhir dapat berubah setelah soal esai diperiksa.\n")
            f.write("--------------------------------------------------\n")
            f.write("KETERANGAN LEMBAR ESAI:\n")
            f.write("Harap tulis caranya di kertas selembar/coret coretan jika tidak maka nilai yang anda dapatkan setengah dari nilai seharusnya!\n")
            f.write("--------------------------------------------------\n")
            f.write("DETAIL JAWABAN PILIHAN GANDA:\n")
            for d in details:
                status_symbol = "BENAR" if d['is_correct'] else "SALAH"
                f.write(f"Soal {d['index']}: {status_symbol}\n")
                f.write(f"  Jawaban Siswa: {d['student_answer'] or '-'}\n")
                f.write(f"  Jawaban Kunci: {d['correct_answer']}\n\n")
                
            if essay_details:
                f.write("--------------------------------------------------\n")
                f.write("DETAIL JAWABAN ESAI (TIDAK DINILAI OTOMATIS):\n")
                for ed in essay_details:
                    f.write(f"Soal Esai {ed['index']}: {ed['question']}\n")
                    f.write(f"  Jawaban Siswa: {ed['student_answer'] or '(Kosong)'}\n\n")
    except Exception as e:
        print(f"[Local Save Warning] Failed writing local result txt file: {e}")

            
    try:
        html_path = os.path.join(result_dir, 'hasil_ujian.html')
        
        essay_rows_html = ""
        if essay_details:
            essay_rows_html = "<h3>Lembar Jawaban Esai:</h3>"
            for ed in essay_details:
                essay_rows_html += f"""
                <div class="detail-item" style="border-left: 5px solid #fd7e14;">
                    <div class="question-txt">Soal Esai {ed['index']}: {ed['question']}</div>
                    <div style="font-size:14px; margin-top:8px;">
                        <strong>Jawaban Siswa:</strong>
                        <div style="background:#fff; border:1px solid #ddd; padding:10px; border-radius:6px; margin-top:5px; white-space:pre-wrap;">{ed['student_answer'] or '<i>Tidak diisi</i>'}</div>
                    </div>
                </div>
                """
        
        report_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Laporan Hasil Ujian - {nama_siswa}</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; padding: 20px; background: #f4f6f9; color: #333; }}
        .report-card {{ background: white; padding: 30px; border-radius: 12px; box-shadow: 0 4px 15px rgba(0,0,0,0.05); max-width: 800px; margin: auto; }}
        .header {{ text-align: center; border-bottom: 2px solid #007bff; padding-bottom: 20px; margin-bottom: 30px; }}
        .school-title {{ font-size: 24px; font-weight: bold; color: #0c2340; text-transform: uppercase; margin: 5px 0; }}
        .report-title {{ font-size: 18px; color: #007bff; margin: 5px 0; }}
        .meta-table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
        .meta-table td {{ padding: 8px 12px; border: 1px dashed #ddd; font-size: 15px; }}
        .meta-label {{ font-weight: bold; width: 150px; background: #f8f9fa; }}
        .score-box {{ text-align: center; margin: 20px 0; padding: 20px; background: #e8f0fe; border-radius: 8px; border: 1px solid #b3d1ff; }}
        .score-val {{ font-size: 48px; font-weight: bold; color: #1a73e8; }}
        .alert-warning-essay {{ background: #fff3cd; color: #856404; border: 1px solid #ffeeba; padding: 15px; border-radius: 8px; font-size: 14px; margin-bottom: 25px; line-height: 1.5; }}
        .detail-item {{ background: #fafafa; border: 1px solid #eee; border-radius: 8px; padding: 15px; margin-bottom: 15px; }}
        .detail-item.correct {{ border-left: 5px solid #28a745; }}
        .detail-item.incorrect {{ border-left: 5px solid #dc3545; }}
        .question-txt {{ font-weight: 500; font-size: 16px; margin-bottom: 10px; }}
        .answer-comparison {{ display: flex; gap: 20px; font-size: 14px; margin-top: 10px; }}
        .ans-badge {{ padding: 4px 10px; border-radius: 4px; font-weight: 500; }}
        .badge-student {{ background: #e9ecef; border: 1px solid #ced4da; }}
        .badge-correct {{ background: #d4edda; border: 1px solid #c3e6cb; color: #155724; }}
        .status-tag {{ font-weight: bold; float: right; font-size: 14px; text-transform: uppercase; }}
        .status-correct {{ color: #28a745; }}
        .status-incorrect {{ color: #dc3545; }}
        .math-fraction {{ display: inline-block; vertical-align: middle; text-align: center; padding: 0 4px; }}
        .math-num {{ display: block; border-bottom: 1px solid #333; padding: 0 2px; }}
        .math-den {{ display: block; padding: 0 2px; }}
    </style>
</head>
<body>
    <div class="report-card">
        <div class="header">
            <div class="school-title">SMK Budi Murni 2</div>
            <div class="report-title">Laporan Lembar Hasil Ujian Matematika</div>
        </div>
        
        <table class="meta-table">
            <tr>
                <td class="meta-label">Nama Siswa</td>
                <td>{nama_siswa}</td>
                <td class="meta-label">Kelas</td>
                <td>{kelas}</td>
            </tr>
            <tr>
                <td class="meta-label">Materi</td>
                <td>{materi}</td>
                <td class="meta-label">Jurusan</td>
                <td>{jurusan}</td>
            </tr>
            <tr>
                <td class="meta-label">Keluar Halaman</td>
                <td style="color:#dc3545; font-weight:bold;">{lock_count} kali</td>
                <td class="meta-label">Waktu Selesai</td>
                <td>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</td>
            </tr>
        </table>

        <div class="alert-warning-essay">
            <strong>KETENTUAN LEMBAR ESAI:</strong><br>
            Harap tulis caranya di kertas selembar/coret coretan jika tidak maka nilai yang anda dapatkan setengah dari nilai seharusnya!
        </div>
        
        <div class="score-box">
            <div>NILAI PILIHAN GANDA (MAKSIMAL 40):</div>
            <div class="score-val">{score} / 40</div>
            <div style="font-size:14px; color:#555;">Keterangan: Benar {correct_count} dari {total_count} soal PG (Bobot PG: 40%, Bobot Essai: 60%)</div>
            <div style="font-size:14px; color:#c2410c; font-weight:bold; margin-top:10px; padding:10px; background:#fff7ed; border:1px solid #ffedd5; border-radius:6px;">
                Catatan: Ini adalah nilai sementara Pilihan Ganda. Nilai akhir dapat berubah setelah soal esai diperiksa.
            </div>
        </div>
        
        <h3>Analisis Lembar Jawaban Pilihan Ganda:</h3>
"""
        
        for d in details:
            status_cls = "correct" if d['is_correct'] else "incorrect"
            status_lbl = "BENAR" if d['is_correct'] else "SALAH"
            status_lbl_cls = "status-correct" if d['is_correct'] else "status-incorrect"
            
            student_ans_val = d['student_answer']
            student_ans_text = f"{student_ans_val}. {d['choices'].get(student_ans_val, '')}" if student_ans_val else "-"
            correct_ans_text = f"{d['correct_answer']}. {d['choices'].get(d['correct_answer'], '')}"
            
            report_html += f"""
            <div class="detail-item {status_cls}">
                <span class="status-tag {status_lbl_cls}">{status_lbl}</span>
                <div class="question-txt">Soal {d['index']}: {d['question']}</div>
                <div class="answer-comparison">
                    <div>Jawaban Siswa: <span class="ans-badge badge-student">{student_ans_text}</span></div>
                    <div>Jawaban Kunci: <span class="ans-badge badge-correct">{correct_ans_text}</span></div>
                </div>
            </div>
            """
            
        report_html += essay_rows_html
        report_html += """
        </div>
</body>
</html>
"""
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(report_html)
    except Exception as e:
        print(f"[Local Save Warning] Failed writing local HTML report: {e}")
        
    return redirect(url_for('selesai'))

@app.route('/selesai')
def selesai():
    if 'siswa' not in session:
        return redirect(url_for('login'))
        
    score = session.get('score', 0.0)
    kelas = session.get('kelas', 'XI')
    materi = session.get('materi', '')
    jurusan = session.get('jurusan', '')
    nama_siswa = session.get('siswa', '')
    
    questions = load_questions(kelas, materi)
    correct_answers = load_answers(kelas, materi)
    essay_questions = load_essay_questions(kelas, materi)
    
    nama_clean = session['nama_clean']
    student_ans = {}
    student_essay_ans = {}
    lock_count = 0
    if nama_clean in ACTIVE_STUDENTS:
        student_ans = ACTIVE_STUDENTS[nama_clean].get('answers', {})
        student_essay_ans = ACTIVE_STUDENTS[nama_clean].get('essay_answers', {})
        lock_count = ACTIVE_STUDENTS[nama_clean].get('lock_count', 0)
    else:
        student_essay_ans = session.get('essay_answers', {})
        lock_count = session.get('lock_count', 0)
        
    details = []
    for q in questions:
        idx = q['index']
        ans = student_ans.get(idx, '')
        expected = correct_answers.get(idx, '')
        details.append({
            'index': idx,
            'question': q['question'],
            'student_answer': ans,
            'correct_answer': expected,
            'is_correct': (ans == expected)
        })
        
    essay_details = []
    for eq in essay_questions:
        idx = eq['index']
        ans = student_essay_ans.get(idx, '')
        essay_details.append({
            'index': idx,
            'question': eq['question'],
            'student_answer': ans
        })
        
    if not kelas.lower().startswith('kelas'):
        folder_kelas = f"Kelas {kelas}"
    else:
        folder_kelas = kelas[0].upper() + kelas[1:]
        
    if materi:
        result_path = f"hasil ujian/{folder_kelas}/{materi}/{jurusan}/{nama_siswa}/"
    else:
        result_path = f"hasil ujian/{folder_kelas}/{jurusan}/{nama_siswa}/"
        
    return render_template('selesai.html', score=score, details=details, essay_details=essay_details, result_path=result_path, materi=materi, lock_count=lock_count)


@app.route('/logout')
def logout():
    session.pop('siswa', None)
    session.pop('kelas', None)
    session.pop('jurusan', None)
    session.pop('nama_clean', None)
    session.pop('status', None)
    session.pop('score', None)
    return redirect(url_for('login'))


# PROCTOR MODULE (SUPERVISOR LOGINS & STATE CONTROLS)

@app.route('/pengawas/login', methods=['GET', 'POST'])
def proctor_login():
    if 'proctor' in session:
        return redirect(url_for('proctor_dashboard'))
        
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        
        if username == 'pengawas1' and password == 'budimurnijaya':
            session['proctor'] = 'active'
            return redirect(url_for('proctor_dashboard'))
        else:
            return render_template('pengawas.html', is_login=True, error='Log masuk salah!')
            
    return render_template('pengawas.html', is_login=True)


@app.route('/pengawas')
def proctor_dashboard():
    if 'proctor' not in session:
        return redirect(url_for('proctor_login'))
        
    return render_template('pengawas.html', is_login=False, 
                           initial_token=CONFIG['INITIAL_TOKEN'], 
                           students=list(ACTIVE_STUDENTS.values()))


@app.route('/pengawas/logout')
def proctor_logout():
    session.pop('proctor', None)
    return redirect(url_for('proctor_login'))


@app.route('/pengawas/api/students')
def proctor_api_students():
    if 'proctor' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    # Return active student list as JSON (allows refreshing the proctor page dynamically)
    return jsonify(list(ACTIVE_STUDENTS.values()))


@app.route('/pengawas/api/unlock/<nama>', methods=['POST'])
def proctor_api_unlock(nama):
    if 'proctor' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    nama_clean = clean_filename(nama)
    if nama_clean in ACTIVE_STUDENTS:
        ACTIVE_STUDENTS[nama_clean]['status'] = 'aktif'
        ACTIVE_STUDENTS[nama_clean]['token_baru'] = ''
        ACTIVE_STUDENTS[nama_clean]['lock_time'] = None
        
        # Remove file
        filepath = f"data_token/{nama_clean}.txt"
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception:
                pass
        return jsonify({'status': 'success'})
        
    return jsonify({'error': 'Student not found'}), 404


@app.route('/pengawas/api/reset/<nama>', methods=['POST'])
def proctor_api_reset(nama):
    if 'proctor' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    nama_clean = clean_filename(nama)
    if nama_clean in ACTIVE_STUDENTS:
        # Delete token file if exists
        filepath = f"data_token/{nama_clean}.txt"
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception:
                pass
                
        # Completely remove/reset student records so they can log back in
        del ACTIVE_STUDENTS[nama_clean]
        return jsonify({'status': 'success'})
        
    return jsonify({'error': 'Student not found'}), 404


@app.route('/pengawas/api/kick/<nama>', methods=['POST'])
def proctor_api_kick(nama):
    if 'proctor' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    nama_clean = clean_filename(nama)
    if nama_clean in ACTIVE_STUDENTS:
        ACTIVE_STUDENTS[nama_clean]['status'] = 'terkunci'
        ACTIVE_STUDENTS[nama_clean]['lock_count'] = ACTIVE_STUDENTS[nama_clean].get('lock_count', 0) + 1
        
        token_baru = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
        ACTIVE_STUDENTS[nama_clean]['token_baru'] = token_baru
        ACTIVE_STUDENTS[nama_clean]['lock_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        try:
            token_dir = os.path.join(BASE_DIR, 'data_token')
            os.makedirs(token_dir, exist_ok=True)
            filename = os.path.join(token_dir, f"{nama_clean}.txt")
            
            content = (
                "=============================\n"
                "TOKEN BARU SISWA (DIKELUARKAN PENGAWAS)\n"
                "=============================\n"
                f"Nama        : {ACTIVE_STUDENTS[nama_clean]['nama']}\n"
                f"Kelas       : {ACTIVE_STUDENTS[nama_clean]['kelas']}\n"
                f"Materi      : {ACTIVE_STUDENTS[nama_clean].get('materi', '-')}\n"
                f"Jurusan     : {ACTIVE_STUDENTS[nama_clean]['jurusan']}\n"
                f"Token       : {token_baru}\n"
                f"Waktu       : {ACTIVE_STUDENTS[nama_clean]['lock_time']}\n"
                f"Catatan     : Dikeluarkan langsung oleh Pengawas Ujian\n"
            )
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
        except Exception as e:
            print(f"[Kick Token File Save Warning] {e}")
            
        return jsonify({'status': 'success', 'token_baru': token_baru})
        
    return jsonify({'error': 'Student not found'}), 404


@app.route('/pengawas/api/update_token', methods=['POST'])
def proctor_api_update_token():
    if 'proctor' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    new_token = request.form.get('new_token', '').strip().upper()
    if not new_token:
        return jsonify({'error': 'Token invalid'}), 400
        
    CONFIG['INITIAL_TOKEN'] = new_token
    return jsonify({'status': 'success', 'new_token': new_token})


@app.route('/api/upload_soal', methods=['POST'])
def api_upload_soal():
    """API endpoint to receive incoming docx question packages from Server Guru."""
    kelas = request.form.get('kelas', '').strip()
    materi = request.form.get('materi', '').strip()
    jurusan = request.form.get('jurusan', 'Semua Jurusan').strip()

    if not kelas or not materi:
        return jsonify({'status': 'error', 'message': 'Data kelas dan materi wajib diisi'}), 400

    base_soal_dir = get_soal_base_dir()
    target_dir = os.path.join(base_soal_dir, kelas, materi)
    try:
        os.makedirs(target_dir, exist_ok=True)
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

    saved_files = []
    for key in ['file_pg', 'file_kunci_pg', 'file_essay', 'file_kunci_essay']:
        if key in request.files:
            file = request.files[key]
            if file and file.filename:
                file_path = os.path.join(target_dir, file.filename)
                file.save(file_path)
                saved_files.append(file.filename)

    meta = {
        'materi': materi,
        'kelas': kelas,
        'jurusan': jurusan,
        'uploaded_by': 'Server Guru Remote',
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    with open(os.path.join(target_dir, 'metadata.json'), 'w', encoding='utf-8') as mf:
        json.dump(meta, mf, indent=2)

    return jsonify({'status': 'success', 'message': f'Soal {materi} ({kelas}) berhasil diterima', 'files': saved_files}), 200


def parse_hasil_txt_file(txt_path):
    info = {}
    if not os.path.exists(txt_path):
        return info
    try:
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if ':' in line:
                    parts = line.split(':', 1)
                    k_strip = parts[0].strip().lower()
                    v_strip = parts[1].strip()
                    if 'nama' in k_strip:
                        info['nama'] = v_strip
                    elif 'kelas' in k_strip:
                        info['kelas'] = v_strip
                    elif 'materi' in k_strip:
                        info['materi'] = v_strip
                    elif 'jurusan' in k_strip:
                        info['jurusan'] = v_strip
                    elif 'tanggal' in k_strip or 'waktu' in k_strip:
                        info['tanggal'] = v_strip
                    elif 'skor pg' in k_strip or 'nilai pg' in k_strip:
                        info['skor_pg'] = v_strip.split('/')[0].strip()
    except Exception:
        pass
    return info

@app.route('/api/get_student_results')
def api_get_student_results():
    """API endpoint to scan local hasil ujian directory and return results JSON to Portal Guru."""
    candidate_dirs = [
        os.path.join(BASE_DIR, 'hasil ujian'),
        os.path.join(os.path.dirname(BASE_DIR), 'hasil ujian'),
        os.path.join(os.getcwd(), 'hasil ujian')
    ]
    hasil_dir = next((d for d in candidate_dirs if os.path.exists(d) and os.path.isdir(d)), os.path.join(BASE_DIR, 'hasil ujian'))
    
    results = []
    classes_set = set()

    if os.path.exists(hasil_dir):
        visited_dirs = set()
        for root, dirs, files in os.walk(hasil_dir):
            if os.path.abspath(root) == os.path.abspath(hasil_dir) or 'readme' in root.lower():
                continue
            
            target_files = [f for f in files if f in ['hasil.txt', 'hasil_ujian.html'] or (f.endswith('.txt') and 'readme' not in f.lower()) or f.endswith('.html')]
            if not target_files or root in visited_dirs:
                continue

            visited_dirs.add(root)
            txt_file = next((f for f in files if f == 'hasil.txt'), None) or next((f for f in files if f.endswith('.txt')), None)
            html_file = next((f for f in files if f == 'hasil_ujian.html'), None) or next((f for f in files if f.endswith('.html')), None)

            main_file = txt_file or html_file
            if not main_file:
                continue

            file_path = os.path.join(root, main_file)
            rel_path = os.path.relpath(file_path, hasil_dir).replace('\\', '/')
            parts = rel_path.split('/')

            parsed = parse_hasil_txt_file(os.path.join(root, txt_file)) if txt_file else {}

            folder_kelas = parts[0] if len(parts) >= 1 else 'Umum'
            folder_materi = parts[1] if len(parts) >= 4 else (parts[1] if len(parts) == 3 else 'Umum')
            folder_jurusan = parts[2] if len(parts) >= 4 else (parts[1] if len(parts) == 3 else 'Semua Jurusan')
            folder_student = parts[3] if len(parts) >= 4 else (parts[2] if len(parts) == 3 else parts[-2] if len(parts) >= 2 else 'Siswa')

            raw_kelas = parsed.get('kelas') or folder_kelas
            norm_kelas = raw_kelas if raw_kelas.startswith('Kelas') else f"Kelas {raw_kelas}"
            materi = parsed.get('materi') or folder_materi
            jurusan = parsed.get('jurusan') or folder_jurusan
            student_name = parsed.get('nama') or folder_student.replace('_', ' ').title()
            mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
            date_str = parsed.get('tanggal') or mod_time
            skor_pg = parsed.get('skor_pg') or '-'

            view_rel = rel_path
            if html_file:
                view_rel = os.path.relpath(os.path.join(root, html_file), hasil_dir).replace('\\', '/')

            classes_set.add(norm_kelas)
            results.append({
                'filename': main_file,
                'rel_path': rel_path,
                'view_rel_path': view_rel,
                'student_name': student_name,
                'kelas': norm_kelas,
                'materi': materi,
                'jurusan': jurusan,
                'skor_pg': skor_pg,
                'date': date_str
            })

    results.sort(key=lambda x: x['date'], reverse=True)
    return jsonify({
        'status': 'success',
        'results': results,
        'classes': sorted(list(classes_set))
    })

@app.route('/api/get_soal_list')
def api_get_soal_list():
    """API endpoint returning scanned questions to Portal Guru."""
    base_dir = get_soal_base_dir()
    materials = []
    if os.path.exists(base_dir):
        for class_folder in sorted(os.listdir(base_dir)):
            class_path = os.path.join(base_dir, class_folder)
            if os.path.isdir(class_path):
                norm_class = class_folder.replace('Kelas ', '').replace('kelas ', '').strip()
                norm_k = class_folder if class_folder.startswith('Kelas') else f"Kelas {class_folder}"
                for mat_name in sorted(os.listdir(class_path)):
                    mat_path = os.path.join(class_path, mat_name)
                    if os.path.isdir(mat_path):
                        files = os.listdir(mat_path)
                        meta_path = os.path.join(mat_path, 'metadata.json')
                        meta = {}
                        if os.path.exists(meta_path):
                            try:
                                with open(meta_path, 'r', encoding='utf-8') as mf:
                                    meta = json.load(mf)
                            except Exception:
                                pass
                        materials.append({
                            'kelas_raw': norm_k,
                            'kelas': norm_class,
                            'materi': mat_name,
                            'jurusan': meta.get('jurusan', 'Semua Jurusan'),
                            'uploaded_by': meta.get('uploaded_by', 'Server Ujian Remote'),
                            'timestamp': meta.get('timestamp', ''),
                            'has_pg': any(f.endswith('.docx') and 'kunci' not in f.lower() and 'essay' not in f.lower() for f in files),
                            'has_key': any(f.endswith('.docx') and 'kunci' in f.lower() and 'essay' not in f.lower() for f in files),
                            'has_essay': any(f.endswith('.docx') and 'essay' in f.lower() for f in files),
                            'files': files
                        })
    return jsonify({'status': 'success', 'materials': materials})

@app.route('/view-hasil/<path:filepath>')
def api_view_hasil(filepath):
    """API endpoint to view or serve student exam result file/HTML."""
    candidate_dirs = [
        os.path.join(BASE_DIR, 'hasil ujian'),
        os.path.join(os.path.dirname(BASE_DIR), 'hasil ujian'),
        os.path.join(os.getcwd(), 'hasil ujian')
    ]
    hasil_dir = next((d for d in candidate_dirs if os.path.exists(d) and os.path.isdir(d)), os.path.join(BASE_DIR, 'hasil ujian'))
    clean_fp = filepath.replace('\\', '/')
    full_path = os.path.abspath(os.path.join(hasil_dir, clean_fp))

    # Check if html exists in same folder
    folder_dir = os.path.dirname(full_path)
    html_candidate = os.path.join(folder_dir, 'hasil_ujian.html')
    if os.path.exists(html_candidate):
        return send_file(html_candidate)

    if os.path.exists(full_path) and full_path.startswith(os.path.abspath(hasil_dir)):
        return send_file(full_path)

@app.route('/api/delete_soal', methods=['POST'])
def api_delete_soal():
    """API endpoint to delete a question folder remotely from Server Guru."""
    kelas = request.form.get('kelas', '').strip()
    materi = request.form.get('materi', '').strip()

    if not kelas or not materi:
        return jsonify({'status': 'error', 'message': 'Kelas dan materi wajib diisi'}), 400

    base_dir = get_soal_base_dir()
    norm_k = kelas.replace('Kelas ', '').replace('kelas ', '').strip()
    candidate_paths = [
        os.path.join(base_dir, kelas, materi),
        os.path.join(base_dir, f"Kelas {norm_k}", materi),
        os.path.join(base_dir, norm_k, materi)
    ]

    target_dir = next((p for p in candidate_paths if os.path.exists(p) and os.path.isdir(p)), None)
    if not target_dir:
        return jsonify({'status': 'error', 'message': 'Folder materi tidak ditemukan di server siswa'}), 404

    try:
        shutil.rmtree(target_dir)
        return jsonify({'status': 'success', 'message': f'Materi {materi} berhasil dihapus dari server siswa'}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500






if __name__ == '__main__':
    try:
        from waitress import serve
        print("==================================================")
        print(" Server Ujian SMK Budi Murni 2 Berjalan!")
        print(" Menggunakan Waitress Multi-Threaded WSGI Server")
        print(" Akses Lokal  : http://localhost:5000")
        print(" Akses LAN/WiFi: http://<IP-Komputer-Server>:5000")
        print(" Threads      : 32 concurrent workers")
        print("==================================================")
        serve(app, host='0.0.0.0', port=5000, threads=32)
    except Exception as e:
        print(f"[Server Warning] Waitress error ({e}), menggunakan server bawaan Flask...")
        app.run(host='0.0.0.0', port=5000, debug=True)

